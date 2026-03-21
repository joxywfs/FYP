import os
import re
from typing import List, Dict, Optional, Union
from dataclasses import dataclass
import PyPDF2
import pdfplumber
import docx
from bs4 import BeautifulSoup
import hashlib
import json
from pathlib import Path

@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document."""
    content: str
    metadata: Dict
    chunk_id: str
    source_id: str
    
    def to_dict(self) -> Dict:
        """Convert chunk to dictionary format."""
        return {
            "content": self.content,
            "metadata": self.metadata,
            "chunk_id": self.chunk_id,
            "source_id": self.source_id
        }

class DocumentProcessor:
    """Handles document processing, text extraction, and chunking."""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.supported_formats = {'.pdf', '.txt', '.docx', '.html', '.md'}
    
    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """Process a document and return chunks."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        # Extract text based on file type
        if extension == '.pdf':
            return self.process_pdf(str(file_path))
        elif extension == '.txt':
            return self.process_text(str(file_path))
        elif extension == '.docx':
            return self.process_docx(str(file_path))
        elif extension == '.html':
            return self.process_html(str(file_path))
        elif extension == '.md':
            return self.process_markdown(str(file_path))
    
    def process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """Extract text from PDF and create chunks."""
        chunks = []
        full_text = ""
        
        try:
            # Try pdfplumber first for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        full_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        except:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        
        # Clean and chunk the text
        full_text = self._clean_text(full_text)
        chunks = self._create_chunks(full_text, file_path)
        
        return chunks
    
    def process_text(self, file_path: str) -> List[DocumentChunk]:
        """Process plain text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        text = self._clean_text(text)
        return self._create_chunks(text, file_path)
    
    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        """Process Word document."""
        doc = docx.Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        text = '\n'.join(full_text)
        text = self._clean_text(text)
        return self._create_chunks(text, file_path)
    
    def process_html(self, file_path: str) -> List[DocumentChunk]:
        """Process HTML file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        text = self._clean_text(text)
        return self._create_chunks(text, file_path)
    
    def process_markdown(self, file_path: str) -> List[DocumentChunk]:
        """Process Markdown file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Basic markdown to text conversion
        # Remove code blocks
        text = re.sub(r'```[\s\S]*?```', '', text)
        # Remove inline code
        text = re.sub(r'`[^`]+`', '', text)
        # Remove links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Remove images
        text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', '', text)
        # Remove headers markdown
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        text = self._clean_text(text)
        return self._create_chunks(text, file_path)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\-.,!?;:\'"()]', ' ', text)
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        return text.strip()
    
    def _create_chunks(self, text: str, source: str) -> List[DocumentChunk]:
        """Create overlapping text chunks with metadata."""
        chunks = []
        words = text.split()
        
        if not words:
            return chunks
        
        # Calculate source file hash for unique identification
        source_hash = hashlib.md5(source.encode()).hexdigest()[:8]
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            
            # Skip empty chunks
            if not chunk_text.strip():
                continue
            
            # Create metadata
            metadata = {
                "source": os.path.basename(source),
                "source_path": source,
                "chunk_index": len(chunks),
                "word_count": len(chunk_words),
                "char_count": len(chunk_text),
                "start_word_index": i,
                "end_word_index": min(i + self.chunk_size, len(words))
            }
            
            # Generate unique chunk ID
            chunk_id = f"{source_hash}_{len(chunks):04d}"
            
            chunk = DocumentChunk(
                content=chunk_text,
                metadata=metadata,
                chunk_id=chunk_id,
                source_id=source_hash
            )
            chunks.append(chunk)
        
        return chunks
    
    def process_directory(self, directory_path: str) -> List[DocumentChunk]:
        """Process all supported documents in a directory."""
        all_chunks = []
        directory = Path(directory_path)
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        # Find all supported files
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    chunks = self.process_document(str(file_path))
                    all_chunks.extend(chunks)
                    print(f"Processed {file_path}: {len(chunks)} chunks")
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
        
        return all_chunks